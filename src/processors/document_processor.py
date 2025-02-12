"""
Procesador de documentos (PDF, XLS, DOC, TXT).
"""
from typing import Dict, Any, Optional
import logging
from pathlib import Path
import PyPDF2
import docx
import pandas as pd
from .base_processor import BaseProcessor

logger = logging.getLogger(__name__)

class DocumentProcessor(BaseProcessor):
    """Procesa documentos en varios formatos."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.xlsx', '.xls', '.docx', '.doc', '.txt'}
    
    async def validate_source(self, source: str) -> bool:
        """Valida si el archivo es compatible."""
        return Path(source).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    async def extract_content(self, source: str) -> Dict[str, Any]:
        """Extrae contenido del documento."""
        try:
            file_path = Path(source)
            extension = file_path.suffix.lower()
            
            if extension in {'.pdf'}:
                return await self._process_pdf(file_path)
            elif extension in {'.xlsx', '.xls'}:
                return await self._process_excel(file_path)
            elif extension in {'.docx', '.doc'}:
                return await self._process_word(file_path)
            elif extension in {'.txt'}:
                return await self._process_text(file_path)
            else:
                raise ValueError(f"Formato no soportado: {extension}")
                
        except Exception as e:
            logger.error(f"Error procesando documento {source}: {str(e)}")
            raise
    
    async def process_content(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Procesa el contenido extraído."""
        try:
            # TODO: Implementar procesamiento específico según tipo de documento
            return {
                "type": content["type"],
                "text": content["text"],
                "metadata": content["metadata"],
                "structured_data": content.get("structured_data", {})
            }
        except Exception as e:
            logger.error(f"Error procesando contenido: {str(e)}")
            raise
    
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo PDF."""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            
            return {
                "type": "pdf",
                "text": text,
                "metadata": {
                    "pages": len(reader.pages),
                    "title": file_path.stem
                }
            }
    
    async def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo Excel."""
        df = pd.read_excel(file_path)
        return {
            "type": "excel",
            "text": df.to_string(),
            "metadata": {
                "sheets": len(pd.ExcelFile(file_path).sheet_names),
                "rows": len(df),
                "columns": len(df.columns)
            },
            "structured_data": {
                "dataframe": df.to_dict(),
                "columns": df.columns.tolist()
            }
        }
    
    async def _process_word(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo Word."""
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return {
            "type": "word",
            "text": text,
            "metadata": {
                "paragraphs": len(doc.paragraphs),
                "title": file_path.stem
            }
        }
    
    async def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Procesa archivo de texto."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            
        return {
            "type": "text",
            "text": text,
            "metadata": {
                "size": file_path.stat().st_size,
                "title": file_path.stem
            }
        }
